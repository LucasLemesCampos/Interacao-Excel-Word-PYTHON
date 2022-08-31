from cgitb import text
from msilib.schema import tables
from docx import Document
import pandas as pd
tabela = pd.read_excel("BasedadosEtiqueta.xlsx")
documento = Document('EtiquetaPADRAO.docx')


contador = 0
for endec in tabela['Endereço']:
    documento = Document('EtiquetaPADRAO.docx')
    for paragrafo in documento.paragraphs:  
        paragrafo.text = paragrafo.text.replace('Endereço!',str(endec))
        documento.save('geradorPython - {0} .docx'.format(contador))
    contador = contador + 1 

contador = 0
for cepp in tabela['CEP']:
    documento = Document('geradorPython - {0} .docx'.format(contador))
    for paragrafo in documento.paragraphs:  
        paragrafo.text = paragrafo.text.replace('YYYY',str(cepp))
        documento.save('geradorPython - {0} .docx'.format(contador))
    contador = contador + 1 

contador = 0
for cpf in tabela['CPF']:
    documento = Document('geradorPython - {0} .docx'.format(contador))
    for paragrafo in documento.paragraphs:  
        paragrafo.text = paragrafo.text.replace('XXXX',str(cpf))
        documento.save('geradorPython - {0} .docx'.format(contador))
    contador = contador + 1 

contador = 0
for pessoa in tabela['user']:
    documento = Document('geradorPython - {0} .docx'.format(contador))
    for paragrafo in documento.paragraphs:  
        paragrafo.text = paragrafo.text.replace('Nome',pessoa)
        documento.save('geradorPython - {0} .docx'.format(contador))
    contador = contador + 1 


    
